from ..models.transcription import Transcription
from sqlalchemy import or_
from ..extensions import db
from datetime import datetime
from flask import send_file
import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from ..services.ai_services import generate_deroulement, analyze_transcription
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display
from decouple import config
import markdown
import markdown2
from bs4 import BeautifulSoup
from docx.shared import Pt
import requests

def create_transcription(data):
    data['dateSceance'] = datetime.strptime(data['dateSceance'], "%Y-%m-%d").date()
    data['DateRedaction'] = datetime.strptime(data['DateRedaction'], "%Y-%m-%d").date()

    if data.get('DateProchaineReunion'):
        data['DateProchaineReunion'] = datetime.strptime(data['DateProchaineReunion'], "%Y-%m-%d").date()
    else:
        data['DateProchaineReunion'] = None

    data['HeureDebut'] = datetime.strptime(data['HeureDebut'], "%H:%M:%S").time()
    data['HeureFin'] = datetime.strptime(data['HeureFin'], "%H:%M:%S").time()

    transcription = Transcription(**data)
    db.session.add(transcription)
    db.session.commit()
    return transcription

def get_transcription_by_id(transcription_id):
    return Transcription.query.get(transcription_id)

def get_all_transcriptions(user_email=None):
    return Transcription.query.filter_by(user_email=user_email).all() if user_email else Transcription.query.all()

def delete_transcription(transcription_id):
    transcription = Transcription.query.get(transcription_id)
    if not transcription:
        return None
    db.session.delete(transcription)
    db.session.commit()
    return transcription

def update_transcription(transcription_id, data):
    transcription = Transcription.query.get(transcription_id)
    if not transcription:
        return None

    if 'dateSceance' in data:
        data['dateSceance'] = datetime.strptime(data['dateSceance'], "%Y-%m-%d").date()
    if 'DateRedaction' in data:
        data['DateRedaction'] = datetime.strptime(data['DateRedaction'], "%Y-%m-%d").date()
    if 'DateProchaineReunion' in data and data['DateProchaineReunion'] is not None:
        data['DateProchaineReunion'] = datetime.strptime(data['DateProchaineReunion'], "%Y-%m-%d").date()

    if 'HeureDebut' in data:
        data['HeureDebut'] = datetime.strptime(data['HeureDebut'], "%H:%M:%S").time()
    if 'HeureFin' in data:
        data['HeureFin'] = datetime.strptime(data['HeureFin'], "%H:%M:%S").time()

    # Serialize dict or list to JSON string if needed
    if isinstance(data.get('Transcription'), (dict, list)):
        import json
        data['Transcription'] = json.dumps(data['Transcription'], ensure_ascii=False)

    for key, value in data.items():
        setattr(transcription, key, value)

    db.session.commit()
    return transcription

def search_transcriptions(query):
    return Transcription.query.filter(
        or_(
            Transcription.titreSceance.ilike(f"%{query}%"),
            Transcription.President.ilike(f"%{query}%"),
            Transcription.OrdreDuJour.ilike(f"%{query}%"),
            Transcription.Resume.ilike(f"%{query}%"),
            Transcription.PV.ilike(f"%{query}%")
        )
    ).all()

def update_transcription_with_deroulement(transcription_id: int) -> Transcription:
    transcription = Transcription.query.get(transcription_id)
    if not transcription or not transcription.Transcription:
        raise ValueError("Transcription not found or missing text.")

    transcription.Deroulement = generate_deroulement(transcription.Transcription)
    db.session.commit()
    return transcription

def update_transcription_with_analysis(transcription_id: int) -> Transcription:
    transcription = Transcription.query.get(transcription_id)
    if not transcription or not transcription.Transcription:
        raise ValueError("Transcription not found or missing text.")

    transcription.Analyse = analyze_transcription(transcription.Transcription)
    db.session.commit()
    return transcription

# -------------- GET NGROK URL ----------------
def get_ngrok_url():
    return config("NGROK_URL") 

# ----------- Export PV DOCX -----------
def export_transcription_pv_arabe_docx(transcription_id):
    
    transcription = Transcription.query.get(transcription_id)
    if not transcription:
        return None

    doc = Document()

    def add_arabic_paragraph(text, bold=False, font_size=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = font_size
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return p

    

    add_arabic_paragraph("محضر اجتماع", bold=True, font_size=Pt(20))

    fields = [
        (transcription.titreSceance, "عنوان الجلسة"),
        (transcription.dateSceance.strftime('%Y/%m/%d'), "تاريخ الجلسة"),
        (
            f"{transcription.HeureDebut.strftime('%H:%M')} إلى {transcription.HeureFin.strftime('%H:%M')}"
            if transcription.HeureDebut and transcription.HeureFin else "غير متوفر",
            "الساعة"
        ),
        (transcription.President, "الرئيس"),
        (transcription.Secretaire, "الكاتب"),
    ]

    for i in range(len(fields)):
        value, label = fields[i]
        if i==0 and i >= 2:
            add_arabic_paragraph(f"{label} :{value}", font_size=Pt(14))
        else:
            add_arabic_paragraph(f"{value} :{label}", font_size=Pt(14))
        
    

    doc.add_paragraph()
    add_arabic_paragraph(":الأعضاء الحاضرون", bold=True, font_size=Pt(20))
    membres = (transcription.Membres or "").split(',')
    for membre in membres:
        membre = membre.strip()
        if membre:
            add_arabic_paragraph(f" {membre} -", font_size=Pt(14))

    if transcription.Absents:
        doc.add_paragraph()
        add_arabic_paragraph(":الأعضاء الغائبون", bold=True, font_size=Pt(20))
        absents = transcription.Absents.split(',')
        for absent in absents:
            absent = absent.strip()
            if absent:
                add_arabic_paragraph(f" {absent} -", font_size=Pt(14))

    doc.add_paragraph()
    add_arabic_paragraph(":جدول الأعمال", bold=True, font_size=Pt(20))
    add_arabic_paragraph(transcription.OrdreDuJour or "غير متوفر", font_size=Pt(14))

    doc.add_paragraph()
    add_arabic_paragraph(":سير الجلسة", bold=True, font_size=Pt(20))
    # --- Markdown to HTML to DOCX ---
    if transcription.Deroulement:
        deroulement_html = markdown2.markdown(transcription.Deroulement)
        soup = BeautifulSoup(deroulement_html, "html.parser")
        for elem in soup.contents:
            # Remove any element containing 'سير الجلسة'
            if 'سير الجلسة' in elem.get_text():
                continue
            if elem.name == 'p':
                text = elem.get_text().strip()
                if not text:
                    continue  # Skip empty paragraphs
                add_arabic_paragraph(text, font_size=Pt(14))
            elif elem.name == 'ul':
                for li in elem.find_all('li'):
                    li_text = li.get_text().strip()
                    if not li_text:
                        continue
                    add_arabic_paragraph(f"• {li_text}", font_size=Pt(14))
            elif elem.name == 'ol':
                for idx, li in enumerate(elem.find_all('li'), 1):
                    li_text = li.get_text().strip()
                    if not li_text:
                        continue
                    add_arabic_paragraph(f"{idx}. {li_text}", font_size=Pt(14))
            elif elem.name == 'strong' or elem.name == 'b':
                strong_text = elem.get_text().strip()
                if not strong_text:
                    continue
                add_arabic_paragraph(strong_text, bold=True, font_size=Pt(14))
            elif elem.name is None:
                plain_text = str(elem).strip()
                if not plain_text:
                    continue
                add_arabic_paragraph(plain_text, font_size=Pt(14))
    else:
        add_arabic_paragraph("غير متوفر", font_size=Pt(14))

    if transcription.DateProchaineReunion:
        doc.add_paragraph()
        add_arabic_paragraph(":تاريخ الاجتماع المقبل", bold=True, font_size=Pt(20))
        add_arabic_paragraph(transcription.DateProchaineReunion.strftime('%Y/%m/%d'), font_size=Pt(14))

    doc.add_paragraph()
    p_points = doc.add_paragraph("..............  حرر بمدينة .............. في تاريخ  ")
    p_points.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_points = doc.add_paragraph("............................. :الرئيس: ............................الكاتب")
    p_points.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"محضر_{transcription.titreSceance}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ----------- Export PV PDF -----------
def export_transcription_pv_arabe_pdf(transcription_id):
    transcription = Transcription.query.get(transcription_id)
    if not transcription:
        return None

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Register Arabic font
    font_path = "static/fonts/Amiri-Regular.ttf"
    pdfmetrics.registerFont(TTFont("Arabic", font_path))

    def draw_right_aligned_line(text, y, font_size=12):
        reshaped = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped)
        c.setFont("Arabic", font_size)
        text_width = c.stringWidth(bidi_text, "Arabic", font_size)
        c.drawString(width - text_width - 50, y, bidi_text)

    y = height - 50
    line_height = 18

    draw_right_aligned_line("محضر اجتماع", y, font_size=16)
    y -= 2 * line_height

    fields = [
        (transcription.titreSceance, "عنوان الجلسة"),
        (transcription.dateSceance.strftime('%Y/%m/%d'), "تاريخ الجلسة"),
        (f"{transcription.HeureDebut.strftime('%H:%M')} إلى {transcription.HeureFin.strftime('%H:%M')}", "الساعة"),
        (transcription.President, "الرئيس"),
        (transcription.Secretaire, "الكاتب"),
    ]

    for value, label in fields:
        draw_right_aligned_line(f"{value} :{label}", y)
        y -= line_height

    y -= line_height
    draw_right_aligned_line(":الأعضاء الحاضرون", y)
    y -= line_height
    for line in (transcription.Membres or "لا يوجد").split('\n'):
        draw_right_aligned_line(line.strip(), y)
        y -= line_height

    if transcription.Absents:
        y -= line_height
        draw_right_aligned_line(":الأعضاء الغائبون", y)
        y -= line_height
        for line in transcription.Absents.split('\n'):
            draw_right_aligned_line(line.strip(), y)
            y -= line_height

    y -= line_height
    draw_right_aligned_line(":جدول الأعمال", y)
    y -= line_height
    for line in (transcription.OrdreDuJour or "غير متوفر").split('\n'):
        draw_right_aligned_line(line.strip(), y)
        y -= line_height

    y -= line_height
    draw_right_aligned_line(":سير الجلسة", y)
    y -= line_height
    for line in (transcription.Deroulement or "غير متوفر").split('\n'):
        draw_right_aligned_line(line.strip(), y)
        y -= line_height

    if transcription.DateProchaineReunion:
        y -= line_height
        draw_right_aligned_line(":تاريخ الاجتماع المقبل", y)
        y -= line_height
        draw_right_aligned_line(transcription.DateProchaineReunion.strftime('%Y/%m/%d'), y)
        y -= line_height

    y -= 2 * line_height
    c.setFont("Arabic", 12)
    draw_right_aligned_line("..............  حرر بمدينة .............. في تاريخ  ", y)
    y -= line_height
    draw_right_aligned_line("............................. :الرئيس: ............................الكاتب", y)

    c.showPage()
    c.save()
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"محضر_{transcription.titreSceance}.pdf",
        mimetype="application/pdf"
    )

# ----------- Export Analysis DOCX -----------
def export_transcription_analysis_arabe_docx(transcription_id):
    transcription = Transcription.query.get(transcription_id)
    if not transcription or not transcription.Analyse:
        return None

    doc = Document()
    from docx.shared import Pt
    from bs4 import BeautifulSoup
    import markdown2
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc.add_paragraph("تحليل محضر الاجتماع", style='Title').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_arabic_paragraph(text, bold=False, font_size=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = font_size
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        return p

    # --- Markdown to HTML to DOCX ---
    analyse_html = markdown2.markdown(
        transcription.Analyse,
        extras=["biblio", "fenced-code-blocks", "tables", "strike", "cuddled-lists", "footnotes"]
    )
    soup = BeautifulSoup(analyse_html, "html.parser")
    for elem in soup.contents:
        # Skip duplicate/unwanted section headers
        if 'تحليل محضر الاجتماع' in elem.get_text():
            continue
        if elem.name == 'h1' or elem.name == 'h2' or elem.name == 'h3':
            text = elem.get_text().strip()
            if not text:
                continue
            add_arabic_paragraph(text, bold=True, font_size=Pt(18))
        elif elem.name == 'p':
            text = elem.get_text().strip()
            if not text:
                continue
            add_arabic_paragraph(text, font_size=Pt(14))
        elif elem.name == 'ul':
            for li in elem.find_all('li'):
                li_text = li.get_text().strip()
                if not li_text:
                    continue
                add_arabic_paragraph(f"• {li_text}", font_size=Pt(14))
        elif elem.name == 'ol':
            for idx, li in enumerate(elem.find_all('li'), 1):
                li_text = li.get_text().strip()
                if not li_text:
                    continue
                add_arabic_paragraph(f"{idx}. {li_text}", font_size=Pt(14))
        elif elem.name in ('strong', 'b'):
            strong_text = elem.get_text().strip()
            if not strong_text:
                continue
            add_arabic_paragraph(strong_text, bold=True, font_size=Pt(14))
        elif elem.name is None:
            plain_text = str(elem).strip()
            if not plain_text:
                continue
            add_arabic_paragraph(plain_text, font_size=Pt(14))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"تحليل_محضر_{transcription.titreSceance}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ----------- Export Analysis PDF -----------
def export_transcription_analysis_arabe_pdf(transcription_id):
    transcription = Transcription.query.get(transcription_id)
    if not transcription or not transcription.Analyse:
        return None

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    textobject = c.beginText()
    textobject.setFont("Helvetica", 12)
    textobject.setTextOrigin(width - 50, height - 50)
    lines = transcription.Analyse.split('\n')
    for line in lines:
        textobject.textLine(line.rjust(100))  # approximate RTL look
    c.drawText(textobject)
    c.showPage()
    c.save()
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"تحليل_محضر_{transcription.titreSceance}.pdf",
        mimetype="application/pdf"
    )
def export_transcription_analysis_arabe_pdf(transcription_id):
    transcription = Transcription.query.get(transcription_id)
    if not transcription or not transcription.Analyse:
        return None

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Register Arabic font
    font_path = "static/fonts/Amiri-Regular.ttf"
    pdfmetrics.registerFont(TTFont("Arabic", font_path))

    y = height - 50
    line_height = 18
    c.setFont("Arabic", 12)

    for line in transcription.Analyse.split('\n'):
        reshaped = arabic_reshaper.reshape(line)
        bidi_text = get_display(reshaped)
        text_width = c.stringWidth(bidi_text, "Arabic", 12)
        c.drawString(width - text_width - 50, y, bidi_text)
        y -= line_height

    c.showPage()
    c.save()
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"تحليل_محضر_{transcription.titreSceance}.pdf",
        mimetype="application/pdf"
    )

def export_transcription_pv_pdfshift(transcription_id):
    """
    Export the PV as a PDF using PDFShift API (Python requests), with Markdown-to-HTML conversion and correct Arabic rendering.
    """
    import requests
    transcription = Transcription.query.get(transcription_id)
    if not transcription:
        return None

    # Convert Deroulement from Markdown to HTML and wrap for RTL/Arabic font
    deroulement_html = "غير متوفر"
    if transcription.Deroulement:
        deroulement_html = markdown.markdown(transcription.Deroulement)
        deroulement_html = f'<div dir="rtl" style="font-family: Amiri, serif; font-size: 16px;">{deroulement_html}</div>'

    # Build HTML content for the PV
    html_content = f"""
    <html>
    <head>
        <meta charset='utf-8'>
        <title>محضر اجتماع</title>
        <style>
            body {{ direction: rtl; font-family: 'Amiri', serif; font-size: 16px; background: #f9f9f9; color: #222; margin: 0; padding: 30px; }}
            h1 {{ font-family: 'Amiri', serif; font-size: 2.2em; color: #2a4d69; text-align: center; margin-bottom: 0.2em; }}
            h2 {{ font-family: 'Amiri', serif; font-size: 1.5em; color: #4b86b4; text-align: center; margin-top: 0; margin-bottom: 1em; }}
            p {{ margin: 0.5em 0 0.5em 0; line-height: 1.8; }}
            b {{ color: #1b3b5a; }}
            /* Inline styles for section-label and deroulement-block */
        </style>
    </head>
    <body>
    <h1>محضر اجتماع</h1>
    <h2>{transcription.titreSceance}</h2>
    <p><b>تاريخ الجلسة:</b> {transcription.dateSceance.strftime('%Y/%m/%d')}</p>
    <p><b>الساعة:</b> {transcription.HeureDebut.strftime('%H:%M')} إلى {transcription.HeureFin.strftime('%H:%M')}</p>
    <p><b>الرئيس:</b> {transcription.President}</p>
    <p><b>الكاتب:</b> {transcription.Secretaire}</p>
    <p><b>الأعضاء الحاضرون:</b> {transcription.Membres or 'لا يوجد'}</p>
    {f'<p><b>الأعضاء الغائبون:</b> {transcription.Absents}</p>' if transcription.Absents else ''}
    <p><b>جدول الأعمال:</b> {transcription.OrdreDuJour or 'غير متوفر'}</p>
    <span style='font-weight:bold; color:#1b3b5a; margin-top:1.2em; display:block;'>سير الجلسة:</span>
    <div style='background:#fff; border:1px solid #e0e0e0; border-radius:8px; padding:18px 20px; margin:1em 0; box-shadow:0 2px 8px #0001;'>{deroulement_html}</div>
    {f'<p><b>تاريخ الاجتماع المقبل:</b> {transcription.DateProchaineReunion.strftime('%Y/%m/%d')}</p>' if transcription.DateProchaineReunion else ''}
    </body></html>
    """
    api_key = config("PDFSHIT_KEY") 
    response = requests.post(
        'https://api.pdfshift.io/v3/convert/pdf',
        headers={ 'X-API-Key': api_key },
        json={
            "source": html_content,
            "landscape": False,
            "use_print": False
        }
    )
    response.raise_for_status()
    from flask import send_file
    import io
    return send_file(
        io.BytesIO(response.content),
        as_attachment=True,
        download_name=f"محضر_{transcription.titreSceance}.pdf",
        mimetype="application/pdf"
    )

def update_transcription_segments(transcription_id, refined_segments):
    import json
    transcription = Transcription.query.get(transcription_id)
    if not transcription:
        return None
    # Store as JSON string for compatibility
    transcription.Transcription = json.dumps(refined_segments, ensure_ascii=False)
    db.session.commit()
    return transcription

def export_transcription_analysis_arabe_pdfshift(transcription_id):
    """
    Export the analysis as a PDF using PDFShift API, with Markdown-to-HTML conversion and correct Arabic rendering.
    """
    import requests
    transcription = Transcription.query.get(transcription_id)
    if not transcription or not transcription.Analyse:
        return None

    # Convert Analyse from Markdown to HTML and wrap for RTL/Arabic font
    import markdown2
    analyse_html = markdown2.markdown(
        transcription.Analyse,
        extras=["biblio", "fenced-code-blocks", "tables", "strike", "cuddled-lists", "footnotes"]
    )
    analyse_html = f'<div dir="rtl" style="font-family: Amiri, serif; font-size: 16px;">{analyse_html}</div>'

    # Build HTML content for the Analysis
    html_content = f"""
    <html>
    <head>
        <meta charset='utf-8'>
        <title>تحليل محضر الاجتماع</title>
        <style>
            body {{ direction: rtl; font-family: 'Amiri', serif; font-size: 16px; background: #f9f9f9; color: #222; margin: 0; padding: 30px; }}
            h1 {{ font-family: 'Amiri', serif; font-size: 2.2em; color: #2a4d69; text-align: center; margin-bottom: 0.2em; }}
            h2 {{ font-family: 'Amiri', serif; font-size: 1.5em; color: #4b86b4; text-align: center; margin-top: 0; margin-bottom: 1em; }}
            p {{ margin: 0.5em 0 0.5em 0; line-height: 1.8; }}
            b, strong {{ color: #1b3b5a; }}
            ul, ol {{ margin-right: 2em; }}
            li {{ margin-bottom: 0.5em; }}
            table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
            th, td {{ border: 1px solid #aaa; padding: 8px; text-align: right; }}
            code, pre {{ background: #f4f4f4; border-radius: 4px; padding: 2px 6px; font-family: 'Cascadia Mono', 'Consolas', monospace; }}
        </style>
    </head>
    <body>
    <h1>تحليل محضر الاجتماع</h1>
    <h2>{transcription.titreSceance}</h2>
    {analyse_html}
    </body></html>
    """
    api_key = config("PDFSHIT_KEY")
    response = requests.post(
        'https://api.pdfshift.io/v3/convert/pdf',
        headers={ 'X-API-Key': api_key },
        json={
            "source": html_content,
            "landscape": False,
            "use_print": False
        }
    )
    response.raise_for_status()
    from flask import send_file
    import io
    return send_file(
        io.BytesIO(response.content),
        as_attachment=True,
        download_name=f"تحليل_محضر_{transcription.titreSceance}.pdf",
        mimetype="application/pdf"
    )
